"""Extractor de Word (.docx)."""
from __future__ import annotations

from pathlib import Path

from ..models import ExtractedDoc
from .base import ExtractionError, Extractor


class DocxExtractor(Extractor):
    extensions = (".docx",)

    def extract(self, path: Path) -> ExtractedDoc:
        try:
            import docx  # python-docx
        except ImportError as exc:  # pragma: no cover
            raise ExtractionError("Falta dependencia: python-docx") from exc

        document = docx.Document(str(path))

        parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
        # Texto de tablas: cada fila como celdas separadas por " | ".
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        core = document.core_properties
        meta = {
            "format": "docx",
            "author": core.author or None,
            "title": core.title or None,
        }
        return ExtractedDoc(text="\n".join(parts), meta=meta)
